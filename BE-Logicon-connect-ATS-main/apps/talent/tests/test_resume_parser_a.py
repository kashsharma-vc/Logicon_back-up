"""
apps/talent/tests/test_resume_parser_a.py

Phase Resume-Parser-Backend-A — comprehensive test suite.

Tests cover:
  - Text extraction (PDF, DOCX, plain text, unsupported format)
  - LLM parser (no API key → ManualReviewRequired)
  - Validation (identifiers, email, list fields)
  - Normalization (name split, phone, skills)
  - Persistence (candidate update, skills, experience, education, ParsedResume)
  - Orchestration / full pipeline
  - process_resume_task (skip, force, duplicate)
  - queue_resume_processing (duplicate file_hash)
  - ResumeViewSet new actions (status, reprocess, mark-reviewed)
  - intake _link_resume_if_needed queues processing
"""

import tempfile
from decimal import Decimal
from io import BytesIO
from unittest.mock import patch, MagicMock

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from rest_framework.test import APIClient

from apps.access.capabilities import CANDIDATE_READ, CANDIDATE_UPDATE, RESUME_READ, RESUME_UPLOAD
from apps.access.models import AccessRole, UserRoleAssignment
from apps.access.tests.utils import bootstrap_role_permissions
from apps.accounts.models import User
from apps.core.models import Organization, ScopeNode
from apps.jobs.models import JobRole
from apps.talent.models import (
    Candidate, Resume, CandidateSkill,
    ParsedResume, CandidateExperience, CandidateEducation,
    ResumeImportBatch, ResumeImportItem,
)
from apps.talent.resume_parser.exceptions import ManualReviewRequired


# ─── Shared helpers ───────────────────────────────────────────────────────────

def _org(suffix):
    return Organization.objects.create(name=f'Org {suffix}', code=f'rp-{suffix}')


def _scope(org):
    return ScopeNode.objects.create(
        org=org, code=org.code, name=org.code,
        node_type='company', parent=None, depth=0,
        path=org.code, is_active=True,
    )


def _user(username, org):
    u = User.objects.create_user(username=username, password='pass')
    u.org = org
    u.save()
    return u


def _candidate(org, phone='9876543210'):
    return Candidate.objects.create(
        org=org,
        phone=phone, phone_normalized=phone,
        first_name='Unknown', last_name='Unknown',
        source='qr',
    )


def _simple_pdf_bytes():
    """Minimal PDF-like bytes for testing (not a real PDF, but sufficient for mock tests)."""
    return b'%PDF-1.4 fake pdf content for testing purposes only'


def _grant_resume_caps(user, org, scope, caps):
    role = AccessRole.objects.get_or_create(
        org=org, code=f'rp_{user.username}', defaults={'name': f'rp_{user.username}'},
    )[0]
    bootstrap_role_permissions(role, caps)
    UserRoleAssignment.objects.create(user=user, role=role, scope_node=scope)


# ─── 1. Extraction ────────────────────────────────────────────────────────────

class TestExtraction(TestCase):
    """Unit tests for extraction helpers using BytesIO directly."""

    def test_clean_text_removes_control_chars(self):
        from apps.talent.resume_parser.extraction import _clean_text
        raw = "Hello\x00World\x07\nFoo\n\n\nBar"
        cleaned = _clean_text(raw)
        self.assertNotIn('\x00', cleaned)
        self.assertNotIn('\x07', cleaned)
        self.assertIn('Bar', cleaned)

    def test_confidence_long_text(self):
        from apps.talent.resume_parser.extraction import _confidence
        raw = 'a' * 1000
        cleaned = 'a' * 950
        c = _confidence(raw, cleaned)
        self.assertGreater(c, 0.5)
        self.assertLessEqual(c, 0.9)

    def test_extract_plain_text_success(self):
        from apps.talent.resume_parser.extraction import _extract_plain_text
        content = b'John Doe\nSenior Security Guard\n10 years experience\nMumbai'
        raw, cleaned, engine, conf = _extract_plain_text(content)
        self.assertEqual(engine, 'plain_text')
        self.assertIn('John Doe', cleaned)
        self.assertGreater(conf, 0)

    def test_extract_plain_text_empty_raises(self):
        from apps.talent.resume_parser.extraction import _extract_plain_text
        with self.assertRaises(ManualReviewRequired):
            _extract_plain_text(b'   \n   ')

    def test_extract_docx_success(self):
        from apps.talent.resume_parser.extraction import _extract_docx
        try:
            from docx import Document
        except ImportError:
            self.skipTest('python-docx not installed')
        buf = BytesIO()
        doc = Document()
        doc.add_paragraph('John Doe - Senior Security Guard')
        doc.add_paragraph('Experience: 10 years in security operations at Mumbai')
        doc.save(buf)
        buf.seek(0)
        raw, cleaned, engine, conf = _extract_docx(buf)
        self.assertEqual(engine, 'python-docx')
        self.assertIn('John Doe', cleaned)

    def test_extract_doc_format_raises(self):
        """Legacy .doc raises ManualReviewRequired."""
        from apps.talent.resume_parser.extraction import extract_text

        org = _org('ext-doc')
        cand = _candidate(org)
        f = SimpleUploadedFile('cv.doc', b'some doc content', content_type='application/msword')
        resume = Resume.objects.create(
            candidate=cand, file=f,
            original_filename='cv.doc',
            content_type='application/msword',
        )
        with self.assertRaises(ManualReviewRequired) as ctx:
            extract_text(resume)
        self.assertIn('doc', str(ctx.exception).lower())


# ─── 2. LLM Parser ───────────────────────────────────────────────────────────

class TestLLMParser(TestCase):

    def test_no_api_key_raises_manual_review(self):
        from apps.talent.resume_parser.llm_parser import parse_resume_text
        with patch.dict('os.environ', {'OPENAI_API_KEY': ''}, clear=False):
            with self.assertRaises(ManualReviewRequired) as ctx:
                parse_resume_text('Some resume text here')
        self.assertIn('OPENAI_API_KEY', str(ctx.exception))

    def test_api_error_raises_manual_review(self):
        from apps.talent.resume_parser.llm_parser import parse_resume_text
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'sk-fake'}, clear=False):
            with patch('openai.OpenAI') as mock_openai:
                mock_client = MagicMock()
                mock_client.chat.completions.create.side_effect = Exception('Connection error')
                mock_openai.return_value = mock_client
                with self.assertRaises(ManualReviewRequired) as ctx:
                    parse_resume_text('Some resume text')
        self.assertIn('API call failed', str(ctx.exception))

    def test_invalid_json_response_raises(self):
        from apps.talent.resume_parser.llm_parser import parse_resume_text
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'sk-fake'}, clear=False):
            with patch('openai.OpenAI') as mock_openai:
                mock_client = MagicMock()
                mock_choice = MagicMock()
                mock_choice.message.content = 'not json at all'
                mock_client.chat.completions.create.return_value = MagicMock(
                    choices=[mock_choice]
                )
                mock_openai.return_value = mock_client
                with self.assertRaises(ManualReviewRequired) as ctx:
                    parse_resume_text('Some resume text')
        self.assertIn('invalid JSON', str(ctx.exception))

    def test_valid_response_returns_dict(self):
        import json
        from apps.talent.resume_parser.llm_parser import parse_resume_text
        payload = json.dumps({'full_name': 'John Doe', 'confidence': 0.9})
        with patch.dict('os.environ', {'OPENAI_API_KEY': 'sk-fake'}, clear=False):
            with patch('openai.OpenAI') as mock_openai:
                mock_client = MagicMock()
                mock_choice = MagicMock()
                mock_choice.message.content = payload
                mock_client.chat.completions.create.return_value = MagicMock(
                    choices=[mock_choice]
                )
                mock_openai.return_value = mock_client
                result = parse_resume_text('Some resume text')
        self.assertEqual(result['full_name'], 'John Doe')


# ─── 3. Validation ───────────────────────────────────────────────────────────

class TestDeterministicParser(TestCase):

    def test_extracts_experience_and_education_entries(self):
        from apps.talent.resume_parser.deterministic_parser import parse_resume_text

        parsed = parse_resume_text(
            """
            Ravi Patil
            Phone: 9876543299
            Email: ravi@example.com

            Experience
            Electrician at Spark Facility Services Jan 2020 - Mar 2024

            Education
            ITI Electrician from Pune Industrial Training Institute 2018
            """
        )

        self.assertEqual(parsed['current_company'], 'Spark Facility Services')
        self.assertEqual(parsed['experience'][0]['job_title'], 'Electrician')
        self.assertEqual(parsed['experience'][0]['company_name'], 'Spark Facility Services')
        self.assertEqual(parsed['experience'][0]['start_date'], '2020-01')
        self.assertEqual(parsed['experience'][0]['end_date'], '2024-03')
        self.assertEqual(parsed['education'][0]['degree'], 'ITI')
        self.assertIn('Pune Industrial Training Institute', parsed['education'][0]['institute'])
        self.assertEqual(parsed['education'][0]['end_year'], 2018)


class TestValidation(TestCase):

    def test_valid_with_all_identifiers(self):
        from apps.talent.resume_parser.validation import validate_parsed_json
        data = {
            'full_name': 'John Doe',
            'email': 'john@example.com',
            'phone': '9876543210',
        }
        errors, missing = validate_parsed_json(data)
        self.assertNotIn('no_identifier', errors)
        self.assertEqual(missing, [])

    def test_no_identifier_sets_error(self):
        from apps.talent.resume_parser.validation import validate_parsed_json
        data = {'full_name': '', 'email': '', 'phone': ''}
        errors, missing = validate_parsed_json(data)
        self.assertIn('no_identifier', errors)
        self.assertIn('full_name', missing)
        self.assertIn('email', missing)
        self.assertIn('phone', missing)

    def test_invalid_email_sets_error(self):
        from apps.talent.resume_parser.validation import validate_parsed_json
        data = {'full_name': 'Jane', 'email': 'not-an-email'}
        errors, missing = validate_parsed_json(data)
        self.assertIn('email', errors)
        self.assertNotIn('no_identifier', errors)

    def test_skills_not_list_auto_corrected(self):
        from apps.talent.resume_parser.validation import validate_parsed_json
        data = {'full_name': 'John', 'skills': 'Security'}
        errors, _ = validate_parsed_json(data)
        self.assertIn('skills', errors)
        self.assertEqual(data['skills'], [])

    def test_only_name_no_identifier_not_triggered(self):
        from apps.talent.resume_parser.validation import validate_parsed_json
        data = {'full_name': 'John Doe'}
        errors, missing = validate_parsed_json(data)
        self.assertNotIn('no_identifier', errors)
        self.assertIn('email', missing)
        self.assertIn('phone', missing)


# ─── 4. Normalization ────────────────────────────────────────────────────────

class TestNormalization(TestCase):

    def test_two_part_name(self):
        from apps.talent.resume_parser.normalization import _split_name
        first, middle, last = _split_name('John Doe')
        self.assertEqual(first, 'John')
        self.assertEqual(middle, '')
        self.assertEqual(last, 'Doe')

    def test_three_part_name(self):
        from apps.talent.resume_parser.normalization import _split_name
        first, middle, last = _split_name('John Kumar Doe')
        self.assertEqual(first, 'John')
        self.assertEqual(middle, 'Kumar')
        self.assertEqual(last, 'Doe')

    def test_single_name(self):
        from apps.talent.resume_parser.normalization import _split_name
        first, middle, last = _split_name('Ramesh')
        self.assertEqual(first, 'Ramesh')
        self.assertEqual(middle, '')
        self.assertEqual(last, '')

    def test_skills_normalized_lowercase(self):
        from apps.talent.resume_parser.normalization import normalize_parsed_json
        data = {
            'full_name': 'John Doe',
            'skills': [{'name': 'Security Operations'}, {'name': 'CCTV Monitoring'}],
        }
        n = normalize_parsed_json(data)
        names = [s['normalized_name'] for s in n['skills']]
        self.assertIn('security operations', names)
        self.assertIn('cctv monitoring', names)

    def test_phone_normalization_strips_91(self):
        from apps.talent.resume_parser.normalization import normalize_parsed_json
        data = {'full_name': 'John', 'phone': '+919876543210'}
        n = normalize_parsed_json(data)
        self.assertEqual(n['phone_normalized'], '9876543210')

    def test_invalid_phone_returns_empty(self):
        from apps.talent.resume_parser.normalization import normalize_parsed_json
        data = {'full_name': 'John', 'phone': '123'}
        n = normalize_parsed_json(data)
        self.assertEqual(n['phone_normalized'], '')

    def test_empty_skill_name_filtered(self):
        from apps.talent.resume_parser.normalization import normalize_parsed_json
        data = {'full_name': 'John', 'skills': [{'name': ''}, {'name': 'Guard'}]}
        n = normalize_parsed_json(data)
        self.assertEqual(len(n['skills']), 1)
        self.assertEqual(n['skills'][0]['name'], 'Guard')


# ─── 5. Persistence ──────────────────────────────────────────────────────────

class TestPersistence(TestCase):

    def setUp(self):
        self.org = _org('pers')
        self.candidate = _candidate(self.org, phone='9876543210')
        f = SimpleUploadedFile('cv.txt', b'resume content', content_type='text/plain')
        self.resume = Resume.objects.create(
            candidate=self.candidate,
            file=f,
            original_filename='cv.txt',
            content_type='text/plain',
            status='validating',
        )

    def _normalized(self, overrides=None):
        base = {
            'full_name': 'John Doe',
            'first_name': 'John',
            'middle_name': '',
            'last_name': 'Doe',
            'email': 'john@example.com',
            'phone': '9876543210',
            'phone_normalized': '9876543210',
            'current_role': 'Senior Guard',
            'current_company': 'Test Corp',
            'current_location': 'Mumbai',
            'total_experience_years': 5.0,
            'skills': [
                {'name': 'Security Operations', 'normalized_name': 'security operations',
                 'proficiency': 'advanced', 'years_experience': 5},
            ],
            'experience': [
                {
                    'job_title': 'Senior Guard',
                    'company_name': 'Test Corp',
                    'start_date': '2020-01',
                    'end_date': None,
                    'is_current': True,
                    'duration_months': 48,
                    'responsibilities': ['Patrolling', 'Access control'],
                }
            ],
            'education': [
                {
                    'degree': 'B.Sc.',
                    'specialization': 'Physics',
                    'institute': 'Mumbai University',
                    'start_year': 2010,
                    'end_year': 2013,
                }
            ],
        }
        if overrides:
            base.update(overrides)
        return base

    def test_creates_parsed_resume(self):
        from apps.talent.resume_parser.persistence import persist_parsed_data
        norm = self._normalized()
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.85)
        self.assertTrue(ParsedResume.objects.filter(resume=self.resume).exists())
        pr = ParsedResume.objects.get(resume=self.resume)
        self.assertEqual(pr.confidence, Decimal('0.85'))

    def test_creates_skills_with_parsed_source(self):
        from apps.talent.resume_parser.persistence import persist_parsed_data
        norm = self._normalized()
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.85)
        skill = CandidateSkill.objects.filter(
            candidate=self.candidate, normalized_skill_name='security operations'
        ).first()
        self.assertIsNotNone(skill)
        self.assertEqual(skill.source, 'parsed')
        self.assertEqual(skill.source_resume_id, self.resume.pk)

    def test_creates_experience(self):
        from apps.talent.resume_parser.persistence import persist_parsed_data
        norm = self._normalized()
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.85)
        exp = CandidateExperience.objects.filter(
            candidate=self.candidate, source_resume=self.resume
        ).first()
        self.assertIsNotNone(exp)
        self.assertEqual(exp.job_title, 'Senior Guard')
        self.assertTrue(exp.is_current)
        self.assertIn('Patrolling', exp.responsibilities)

    def test_creates_education(self):
        from apps.talent.resume_parser.persistence import persist_parsed_data
        norm = self._normalized()
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.85)
        edu = CandidateEducation.objects.filter(
            candidate=self.candidate, source_resume=self.resume
        ).first()
        self.assertIsNotNone(edu)
        self.assertEqual(edu.degree, 'B.Sc.')
        self.assertEqual(edu.end_year, 2013)

    def test_updates_candidate_generic_name(self):
        from apps.talent.resume_parser.persistence import persist_parsed_data
        norm = self._normalized()
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.85)
        self.candidate.refresh_from_db()
        self.assertEqual(self.candidate.first_name, 'John')
        self.assertEqual(self.candidate.last_name, 'Doe')

    def test_does_not_overwrite_manual_skill(self):
        from apps.talent.resume_parser.persistence import persist_parsed_data
        CandidateSkill.objects.create(
            candidate=self.candidate,
            skill_name='Security Operations',
            normalized_skill_name='security operations',
            source='manual',
            proficiency='expert',
        )
        norm = self._normalized()
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.85)
        skill = CandidateSkill.objects.get(
            candidate=self.candidate, normalized_skill_name='security operations', source='manual'
        )
        self.assertEqual(skill.proficiency, 'expert')

    def test_replaces_previous_parsed_skills_on_reprocess(self):
        from apps.talent.resume_parser.persistence import persist_parsed_data
        CandidateSkill.objects.create(
            candidate=self.candidate,
            skill_name='Old Skill',
            normalized_skill_name='old skill',
            source='parsed',
            source_resume=self.resume,
        )
        norm = self._normalized()
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.85)
        self.assertFalse(
            CandidateSkill.objects.filter(normalized_skill_name='old skill').exists()
        )

    def test_upserts_parsed_resume_on_reprocess(self):
        from apps.talent.resume_parser.persistence import persist_parsed_data
        norm = self._normalized()
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.85)
        norm['summary'] = 'Updated summary'
        persist_parsed_data(self.resume, norm, norm, {}, [], 0.90)
        self.assertEqual(ParsedResume.objects.filter(resume=self.resume).count(), 1)
        pr = ParsedResume.objects.get(resume=self.resume)
        self.assertEqual(pr.confidence, Decimal('0.90'))


# ─── 6. Orchestration ────────────────────────────────────────────────────────

_MOCK_PARSED = {
    'full_name': 'Ravi Kumar',
    'email': 'ravi@example.com',
    'phone': '9876543211',
    'summary': 'Security professional.',
    'career_level': 'mid',
    'primary_domain': 'Security',
    'total_experience_years': 4.0,
    'current_company': 'Guard Corp',
    'current_role': 'Guard',
    'current_location': 'Pune',
    'skills': [{'name': 'Patrol', 'proficiency': 'intermediate', 'years_experience': 4}],
    'experience': [{
        'job_title': 'Guard', 'company_name': 'Guard Corp',
        'start_date': '2020-01', 'end_date': None,
        'is_current': True, 'duration_months': 48, 'responsibilities': ['Patrolling'],
    }],
    'education': [{'degree': 'SSC', 'specialization': '', 'institute': 'State Board',
                   'start_year': 2010, 'end_year': 2012}],
    'confidence': 0.88,
}


class TestOrchestration(TestCase):

    def setUp(self):
        self.org = _org('orch')
        self.candidate = _candidate(self.org, phone='9876543211')

    def _make_resume(self, content=None, content_type='text/plain', ext='txt'):
        content = content or (b'John Doe\nSenior Guard\nMumbai\n10 years experience in security field at ABC Corp')
        f = SimpleUploadedFile(f'cv.{ext}', content, content_type=content_type)
        return Resume.objects.create(
            candidate=self.candidate,
            file=f,
            original_filename=f'cv.{ext}',
            content_type=content_type,
            status='uploaded',
        )

    def test_no_api_key_still_indexes_with_deterministic_parser(self):
        """Without OPENAI_API_KEY the pipeline stops at parsing → manual_review."""
        resume = self._make_resume()
        with patch.dict('os.environ', {'OPENAI_API_KEY': ''}, clear=False):
            from apps.talent.resume_parser.orchestration import run_pipeline
            run_pipeline(resume)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'indexed')
        self.assertEqual(resume.parser_engine, 'deterministic_v1')

    def test_short_text_sets_manual_review(self):
        """Text shorter than 50 chars → manual_review."""
        resume = self._make_resume(content=b'Hi')
        from apps.talent.resume_parser.orchestration import run_pipeline
        run_pipeline(resume)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'manual_review')
        self.assertIn('short', resume.manual_review_reason.lower())

    def test_full_pipeline_sets_indexed(self):
        """Mock LLM → full pipeline → indexed."""
        resume = self._make_resume()
        from apps.talent.resume_parser.orchestration import run_pipeline
        run_pipeline(resume)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'indexed')
        self.assertEqual(resume.parser_engine, 'deterministic_v1')
        self.assertTrue(ParsedResume.objects.filter(resume=resume).exists())

    def test_full_pipeline_creates_skills(self):
        resume = self._make_resume()
        from apps.talent.resume_parser.orchestration import run_pipeline
        run_pipeline(resume)
        self.assertTrue(
            CandidateSkill.objects.filter(
                candidate=self.candidate, normalized_skill_name='security'
            ).exists()
        )

    def test_extraction_failure_sets_manual_review(self):
        """If extraction raises ManualReviewRequired → manual_review, not failed."""
        resume = self._make_resume()
        with patch('apps.talent.resume_parser.extraction.extract_text',
                   side_effect=ManualReviewRequired('Cannot open file')):
            from apps.talent.resume_parser.orchestration import run_pipeline
            run_pipeline(resume)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'manual_review')

    def test_unexpected_exception_sets_failed(self):
        """Unexpected exception in the pipeline → failed."""
        resume = self._make_resume()
        with patch('apps.talent.resume_parser.extraction.extract_text',
                   side_effect=RuntimeError('disk full')):
            from apps.talent.resume_parser.orchestration import run_pipeline
            run_pipeline(resume)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'failed')
        self.assertIn('disk full', resume.error_message)

    def test_no_identifier_sets_manual_review(self):
        """LLM returns no name/email/phone → validation fails → manual_review."""
        empty_parsed = {**_MOCK_PARSED, 'full_name': '', 'email': '', 'phone': ''}
        resume = self._make_resume()
        with patch('apps.talent.resume_parser.deterministic_parser.parse_resume_text',
                   return_value=empty_parsed):
            from apps.talent.resume_parser.orchestration import run_pipeline
            run_pipeline(resume)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'manual_review')


# ─── 7. process_resume_task ───────────────────────────────────────────────────

class TestProcessResumeTask(TestCase):

    def setUp(self):
        self.org = _org('task')
        self.candidate = _candidate(self.org, phone='9876543212')

    def _resume(self, status='uploaded', file_hash=''):
        f = SimpleUploadedFile(
            'cv.txt',
            b'John Doe\nSecurity guard with ten years experience in Mumbai site operations',
            content_type='text/plain',
        )
        return Resume.objects.create(
            candidate=self.candidate, file=f,
            original_filename='cv.txt', content_type='text/plain',
            status=status, file_hash=file_hash,
        )

    def test_already_indexed_skipped(self):
        resume = self._resume(status='indexed')
        from apps.talent.tasks import process_resume_task
        process_resume_task(resume.pk)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'indexed')

    def test_force_reprocesses_indexed(self):
        """force=True re-runs the pipeline even when already indexed."""
        resume = self._resume(status='indexed')
        with patch.dict('os.environ', {'OPENAI_API_KEY': ''}, clear=False):
            from apps.talent.tasks import process_resume_task
            process_resume_task(resume.pk, force=True)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'indexed')

    def test_duplicate_file_skipped_always(self):
        resume = self._resume(status='duplicate_file')
        from apps.talent.tasks import process_resume_task
        process_resume_task(resume.pk, force=True)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'duplicate_file')

    def test_nonexistent_resume_id_does_not_raise(self):
        from apps.talent.tasks import process_resume_task
        process_resume_task(99999)

    def test_task_full_pipeline_with_deterministic_parser(self):
        resume = self._resume()
        from apps.talent.tasks import process_resume_task
        process_resume_task(resume.pk)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'indexed')


# ─── 8. queue_resume_processing ──────────────────────────────────────────────

class TestQueueResumeProcessing(TestCase):

    def setUp(self):
        self.org = _org('queue')
        self.candidate = _candidate(self.org, phone='9876543213')

    def _resume(self, file_hash=''):
        f = SimpleUploadedFile('cv.txt', b'test content', content_type='text/plain')
        return Resume.objects.create(
            candidate=self.candidate, file=f,
            original_filename='cv.txt', content_type='text/plain',
            status='uploaded', file_hash=file_hash,
        )

    def test_duplicate_file_hash_sets_duplicate_status(self):
        """If an indexed resume with same file_hash exists → duplicate_file."""
        existing = self._resume(file_hash='abc123hash')
        Resume.objects.filter(pk=existing.pk).update(status='indexed')

        new_resume = self._resume(file_hash='abc123hash')
        with patch('apps.talent.tasks.process_resume_task') as mock_task:
            from apps.talent.services import queue_resume_processing
            queue_resume_processing(new_resume)
        new_resume.refresh_from_db()
        self.assertEqual(new_resume.status, 'duplicate_file')
        mock_task.delay.assert_not_called()

    def test_no_duplicate_schedules_task(self):
        resume = self._resume(file_hash='uniquehash999')
        with patch('apps.talent.tasks.process_resume_task') as mock_task:
            mock_task.delay = MagicMock()
            from apps.talent.services import queue_resume_processing
            queue_resume_processing(resume)
        resume.refresh_from_db()
        self.assertEqual(resume.status, 'extracting')
        mock_task.delay.assert_called_once_with(resume.pk)

    def test_empty_file_hash_no_duplicate_check(self):
        """file_hash='' skips duplicate check and schedules normally."""
        resume = self._resume(file_hash='')
        with patch('apps.talent.tasks.process_resume_task') as mock_task:
            mock_task.delay = MagicMock()
            from apps.talent.services import queue_resume_processing
            queue_resume_processing(resume)
        self.assertEqual(resume.status, 'extracting')
        mock_task.delay.assert_called_once()


# ─── 9. ResumeViewSet new actions ────────────────────────────────────────────

class TestResumeViewSetActions(TestCase):

    def setUp(self):
        self.org = _org('api')
        self.scope = _scope(self.org)
        self.candidate = _candidate(self.org)
        self.user = _user('u_rp_api', self.org)
        _grant_resume_caps(
            self.user, self.org, self.scope,
            [RESUME_READ, RESUME_UPLOAD, CANDIDATE_READ, CANDIDATE_UPDATE],
        )
        f = SimpleUploadedFile('cv.txt', b'test', content_type='text/plain')
        self.resume = Resume.objects.create(
            candidate=self.candidate, file=f,
            original_filename='cv.txt', content_type='text/plain',
            status='manual_review',
            manual_review_reason='No API key',
        )
        self.client = APIClient()
        self.client.force_authenticate(self.user)

    def test_status_endpoint_returns_200(self):
        resp = self.client.get(f'/api/talent/resumes/{self.resume.pk}/status/')
        self.assertEqual(resp.status_code, 200, resp.data)
        self.assertEqual(resp.data['status'], 'manual_review')
        self.assertIn('manual_review_reason', resp.data)
        self.assertIn('error_message', resp.data)

    def test_status_endpoint_requires_resume_read(self):
        other_user = _user('u_rp_norole', self.org)
        c = APIClient()
        c.force_authenticate(other_user)
        resp = c.get(f'/api/talent/resumes/{self.resume.pk}/status/')
        self.assertEqual(resp.status_code, 403)

    def test_reprocess_endpoint_returns_200(self):
        with patch('apps.talent.tasks.process_resume_task') as mock_task:
            mock_task.delay = MagicMock()
            resp = self.client.post(f'/api/talent/resumes/{self.resume.pk}/reprocess/')
        self.assertEqual(resp.status_code, 200, resp.data)
        self.assertIn('scheduled', resp.data['detail'].lower())
        mock_task.delay.assert_called_once_with(self.resume.pk, force=True)

    def test_reprocess_duplicate_file_blocked(self):
        Resume.objects.filter(pk=self.resume.pk).update(status='duplicate_file')
        resp = self.client.post(f'/api/talent/resumes/{self.resume.pk}/reprocess/')
        self.assertEqual(resp.status_code, 400)

    def test_mark_reviewed_sets_indexed(self):
        ParsedResume.objects.create(
            resume=self.resume,
            parsed_json={'full_name': 'John Doe'},
            normalized_json={},
            confidence=Decimal('0.75'),
        )
        resp = self.client.post(f'/api/talent/resumes/{self.resume.pk}/mark-reviewed/')
        self.assertEqual(resp.status_code, 200, resp.data)
        self.resume.refresh_from_db()
        self.assertEqual(self.resume.status, 'indexed')

    def test_mark_reviewed_without_parsed_data_blocked(self):
        resp = self.client.post(f'/api/talent/resumes/{self.resume.pk}/mark-reviewed/')
        self.assertEqual(resp.status_code, 400)

    def test_mark_reviewed_non_manual_review_blocked(self):
        Resume.objects.filter(pk=self.resume.pk).update(status='failed')
        resp = self.client.post(f'/api/talent/resumes/{self.resume.pk}/mark-reviewed/')
        self.assertEqual(resp.status_code, 400)

    def test_bulk_upload_creates_async_batch_and_item_task_tags_role(self):
        role = JobRole.objects.create(
            org=self.org,
            name='Electrician',
            code='electrician-rp',
            skill_category='skilled',
        )
        f = SimpleUploadedFile(
            'electrician.txt',
            (
                b'Ramesh Patil\n'
                b'Phone: 9876543219\n'
                b'Electrician with 6 years experience in Pune maintenance operations'
            ),
            content_type='text/plain',
        )

        with patch('apps.talent.tasks.process_resume_import_item_task.delay') as mock_delay:
            with self.captureOnCommitCallbacks(execute=True):
                resp = self.client.post(
                    '/api/talent/resumes/bulk-upload/',
                    {
                        'target_job_role': role.pk,
                        'source_type': 'bulk_upload',
                        'files': [f],
                    },
                    format='multipart',
                )

        self.assertEqual(resp.status_code, 201, resp.data)
        self.assertEqual(resp.data['status'], 'queued')
        self.assertEqual(resp.data['total_count'], 1)
        self.assertEqual(resp.data['processed_count'], 0)
        batch = ResumeImportBatch.objects.get(pk=resp.data['id'])
        item = ResumeImportItem.objects.get(batch=batch)
        self.assertEqual(item.document_type, 'txt')
        mock_delay.assert_called_once_with(item.pk)

        from apps.talent.tasks import process_resume_import_item_task
        process_resume_import_item_task(item.pk)

        candidate = Candidate.objects.get(phone_normalized='9876543219')
        self.assertEqual(candidate.target_job_role_id, role.pk)
        resume = Resume.objects.get(candidate=candidate)
        self.assertEqual(resume.target_job_role_id, role.pk)
        self.assertEqual(resume.document_type, 'txt')
        self.assertEqual(resume.parser_engine, 'deterministic_v1')
        self.assertEqual(resume.status, 'indexed')
        item.refresh_from_db()
        batch.refresh_from_db()
        self.assertEqual(item.status, 'indexed')
        self.assertEqual(batch.status, 'completed')
        self.assertEqual(batch.success_count, 1)

        detail = self.client.get(f'/api/talent/resumes/import-batches/{batch.pk}/')
        self.assertEqual(detail.status_code, 200, detail.data)
        self.assertEqual(detail.data['processed_count'], 1)

        source_filtered = self.client.get(
            '/api/talent/candidates/',
            {'source_type': 'bulk_upload'},
        )
        self.assertEqual(source_filtered.status_code, 200, source_filtered.data)
        results = source_filtered.data.get('results', source_filtered.data)
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]['id'], candidate.pk)

    def test_excel_import_csv_tags_candidate_to_role(self):
        role = JobRole.objects.create(
            org=self.org,
            name='Plumber',
            code='plumber-rp',
            skill_category='skilled',
        )
        csv_file = SimpleUploadedFile(
            'candidates.csv',
            (
                b'full_name,phone,email,skills\n'
                b'Ganesh More,9876543220,ganesh@example.com,Plumbing;Maintenance\n'
            ),
            content_type='text/csv',
        )

        resp = self.client.post(
            '/api/talent/resumes/excel-import/',
            {
                'target_job_role': role.pk,
                'source_type': 'excel_import',
                'file': csv_file,
            },
            format='multipart',
        )

        self.assertEqual(resp.status_code, 201, resp.data)
        self.assertEqual(resp.data['imported'], 1)
        self.assertEqual(resp.data['document_type'], 'csv')
        candidate = Candidate.objects.get(phone_normalized='9876543220')
        self.assertEqual(candidate.target_job_role_id, role.pk)
        self.assertEqual(candidate.skills.count(), 2)
        batch = ResumeImportBatch.objects.get(pk=resp.data['batch_id'])
        self.assertEqual(batch.document_type, 'csv')
        self.assertTrue(batch.import_file.name.endswith('.csv'))
        item = ResumeImportItem.objects.get(batch=batch)
        self.assertEqual(item.document_type, 'csv')
        self.assertEqual(item.row_number, 2)
        self.assertEqual(item.candidate_id, candidate.pk)

        filtered = self.client.get('/api/talent/candidates/', {'document_type': 'csv'})
        self.assertEqual(filtered.status_code, 200, filtered.data)
        results = filtered.data.get('results', filtered.data)
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]['id'], candidate.pk)
        self.assertEqual(results[0]['latest_document_type'], 'csv')
        self.assertEqual(results[0]['latest_source_type'], 'excel_import')

        source_filtered = self.client.get(
            '/api/talent/candidates/',
            {'source_type': 'excel_import'},
        )
        self.assertEqual(source_filtered.status_code, 200, source_filtered.data)
        results = source_filtered.data.get('results', source_filtered.data)
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]['id'], candidate.pk)


# ─── 10. intake _link_resume_if_needed ───────────────────────────────────────

class TestIntakeLinkResume(TestCase):

    def setUp(self):
        from apps.intake.models import QRCampaign, IntakeSubmission, IntakeDocument
        from apps.sites.models import Client, SiteProfile

        self.org = _org('intake-rp')
        self.candidate = _candidate(self.org, phone='9876543214')

        client = Client.objects.create(org=self.org, name='C', code='c-rp1')
        site = SiteProfile.objects.create(
            org=self.org, client=client, name='S', code='s-rp1', is_active=True,
        )
        campaign = QRCampaign.objects.create(
            org=self.org, site=site, name='Camp', code='camp-rp', is_active=True,
        )
        submission = IntakeSubmission.objects.create(
            campaign=campaign,
            site=site,
            candidate=self.candidate,
            mobile_number='9876543214',
            mobile_number_normalized='9876543214',
        )
        f = SimpleUploadedFile('resume.pdf', b'%PDF-1.4 sample text', content_type='application/pdf')
        self.doc = IntakeDocument.objects.create(
            submission=submission,
            document_type='resume',
            file=f,
            original_filename='resume.pdf',
            content_type='application/pdf',
            size_bytes=20,
        )

    def test_link_resume_creates_resume_and_queues(self):
        with patch('apps.talent.tasks.process_resume_task') as mock_task:
            mock_task.delay = MagicMock()
            from apps.intake.services import _link_resume_if_needed
            _link_resume_if_needed(self.doc, self.candidate)

        resume = Resume.objects.filter(source_intake_document=self.doc).first()
        self.assertIsNotNone(resume)
        self.assertEqual(resume.source_type, 'qr_intake')
        # queue_resume_processing was called, which called delay
        mock_task.delay.assert_called_once_with(resume.pk)

    def test_link_resume_idempotent(self):
        """Calling twice does not create a second Resume."""
        with patch('apps.talent.tasks.process_resume_task') as mock_task:
            mock_task.delay = MagicMock()
            from apps.intake.services import _link_resume_if_needed
            _link_resume_if_needed(self.doc, self.candidate)
            _link_resume_if_needed(self.doc, self.candidate)
        self.assertEqual(
            Resume.objects.filter(source_intake_document=self.doc).count(), 1
        )

    def test_non_resume_doc_ignored(self):
        self.doc.document_type = 'id_proof'
        self.doc.save()
        with patch('apps.talent.tasks.process_resume_task') as mock_task:
            mock_task.delay = MagicMock()
            from apps.intake.services import _link_resume_if_needed
            _link_resume_if_needed(self.doc, self.candidate)
        mock_task.delay.assert_not_called()
